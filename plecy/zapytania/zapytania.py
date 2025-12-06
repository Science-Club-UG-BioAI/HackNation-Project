import sqlite3
import pandas as pd
import openpyxl
from docx import Document
import os

# Ensure the hash directory exists
os.makedirs(r"plecy\hash", exist_ok=True)

conn = sqlite3.connect(r"plecy\hash\moja_baza.db")
cursor = conn.cursor()

with open(r"plecy\Projekt_tabele.sql", "r", encoding= "utf-8") as f:
    sql_script = f.read()

cursor.executescript(sql_script)
conn.commit()

def get_planets(conn):
    return conn.execute("SELECT * FROM PLANETY").fetchall()

def planets():
    return get_planets(conn)


def export_planets_to_excel(conn):
    df = pd.read_sql_query("SELECT * FROM planety", conn)
    df.to_excel(r"plecy\hash\planety.xlsx", index=False)
    conn.close()
    os.remove(r"plecy\hash\moja_baza.db")
    return "planety.xlsx"





def export_planets_to_word(conn):
    doc = Document()
    doc.add_heading("Planety", level=1)
    planets = get_planets(conn)
    for planet in planets:
        doc.add_paragraph(f"Nazwa: {planet[1]}, Masa: {planet[2]}, Promień: {planet[3]}")
    doc.save("planety.docx")
    return "planety.docx"

export_planets_to_excel(conn)